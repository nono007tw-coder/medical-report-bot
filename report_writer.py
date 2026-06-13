from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from styles import configure_document, set_run_font, style_table

PATIENT_GROUP_ORDER = ["1. 血液檢查", "2. 生化檢查", "3. 尿液檢查", "4. 影像檢查"]
HEMATOLOGY_CATEGORIES = {"血液常規檢查", "凝血功能檢查"}
PATIENT_CATEGORY_LABELS = {
    "血液常規檢查": "血球與貧血",
    "凝血功能檢查": "凝血功能",
    "肝膽功能檢查": "肝膽功能",
    "腎功能檢查": "腎臟功能",
    "電解質與酸鹼檢查": "電解質與酸鹼",
    "血糖與糖尿病相關檢查": "血糖與糖尿病",
    "血脂檢查": "血脂",
    "鐵質與貧血相關檢查": "鐵質與貧血",
    "發炎與感染指標": "發炎與感染",
    "尿沉渣": "尿液顯微鏡檢查",
    "超音波檢查": "超音波",
    "一般生化檢查": "一般生化",
}


def _result_text(row):
    return " ".join(part for part in (row["result"], row["unit"]) if part).strip()


def _add_title(doc):
    paragraph = doc.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("檢查報告整理")
    set_run_font(run, size=20, bold=True)


def _add_lab_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    headers = ["中文項目", "英文項目", "結果", "正常值"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for row in rows:
        cells = table.add_row().cells
        values = [row["zh"], row["en"], _result_text(row), row["reference"]]
        for cell, value in zip(cells, values):
            cell.text = value
        if row["flag"] in {"H", "L"}:
            for run in cells[2].paragraphs[0].runs:
                run.bold = True
    style_table(table, [3.2, 6.0, 3.0, 4.9])


def _patient_group(section, category):
    if section == "影像檢查":
        return "4. 影像檢查"
    if section == "驗尿檢查":
        return "3. 尿液檢查"
    if category in HEMATOLOGY_CATEGORIES:
        return "1. 血液檢查"
    return "2. 生化檢查"


def _patient_grouped(grouped):
    result = {group: {} for group in PATIENT_GROUP_ORDER}
    for section, categories in grouped.items():
        for category, rows in categories.items():
            group = _patient_group(section, category)
            label = PATIENT_CATEGORY_LABELS.get(category, category.replace("檢查", ""))
            result[group].setdefault(label, []).extend(rows)
    return {group: categories for group, categories in result.items() if categories}


def write_report(grouped, output_path):
    doc = Document()
    configure_document(doc)
    _add_title(doc)

    for section, categories in _patient_grouped(grouped).items():
        doc.add_paragraph(section, style="Heading 1")
        for category, rows in categories.items():
            doc.add_paragraph(category, style="Heading 2")
            normalized_rows = [
                {
                    **row,
                    "result": row["raw_text"] if row["is_image"] else row["result"],
                    "unit": "" if row["is_image"] else row["unit"],
                    "reference": "" if row["is_image"] else row["reference"],
                }
                for row in rows
            ]
            _add_lab_table(doc, normalized_rows)
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(2)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
