import tempfile
import unittest
from pathlib import Path

from docx import Document

from classifier import classify_items
from parser import parse_text
from report_writer import write_report


class ParserTests(unittest.TestCase):
    def test_irregular_cbc_differential_format(self):
        text = """檢體　　：
(Specimen type)\tBlood
醫囑名稱：
(Medical order)\tDC,WBC,Hb,Plt,
項目\tH/L\t結果\t前次結果\t 單位\t  參考值
WBC\t   \t6600\t(  )\t/uL\t4180 ~ 9380
RBC\t   \t\t(  )\t\t
Hb\tL\t8.7\t(  )\tg/dL\t13.3 ~ 17.4
Plt\t   \t171000\t(  )\t/uL\t145000 ~ 383000
ANC\t   \t4415\t(  )\t/uL\t1890 ~ 6760
DC :%\t\t\t\t\t
Band\t   \t0.0\t(  )\t%\t0 ~ 5
Neu.\t   \t66.9\t(  )\t%\t45.0 ~ 81.2
Lym.\t   \t14.8\t(  )\t%\t13.0 ~ 45.5
Mono.\tH\t10.5\t(  )\t%\t3.7 ~ 8.8
Eos.\tH\t7.6\t(  )\t%\t0.2 ~ 6.6
Baso.\t   \t0.2\t(  )\t%\t0.1 ~ 1.6
"""
        items = parse_text(text)
        grouped = classify_items(items)
        rows = grouped["抽血檢查"]["血液常規檢查"]

        self.assertNotIn("DC", [item.raw_name for item in items])
        self.assertEqual(
            [row["key"] for row in rows],
            [
                "WBC", "Hb", "Platelet", "ANC", "Band",
                "Neutrophil", "Lymphocyte", "Monocyte",
                "Eosinophil", "Basophil",
            ],
        )
        self.assertEqual(rows[1]["result"], "8.7")
        self.assertEqual(rows[1]["flag"], "L")
        self.assertEqual(rows[7]["flag"], "H")
        self.assertEqual(rows[8]["reference"], "0.2 ~ 6.6")

    def test_hospital_tabular_format(self):
        text = """檢體　　：
(Specimen type)\tBlood
醫囑名稱：
(Medical order)\tNA,K,Ca,P,Crea,BUN,CRP,ALT,Bil-T,Alb,
項目\tH/L\t結果\t前次結果\t單位\t參考值
BUN\t H  \t 34  \t(  )\t mg/dL \t 6~20 mg/dL
Na\t L  \t 133  \t(  )\t mmol/L \t 136~145 mmol/L
K\t   \t 4.3 \t(  )\t mmol/L \t 3.5~5.1 mmol/L
Cl\t   \t  \t(  )\t  \t
GLU\t   \t  \t(  )\t  \t
Creat\t H  \t 6.34 \t(  )\t mg/dL \t M:0.7~1.2; F:0.5-0.9 mg/dL
ALT\t   \t 15 \t(  )\t U/L \t M:<41; F:<33 U/L
"""
        items = parse_text(text)

        self.assertEqual(
            [item.raw_name for item in items],
            ["BUN", "Na", "K", "Cl", "GLU", "Creat", "ALT"],
        )
        self.assertEqual(items[0].specimen, "Blood")
        self.assertEqual(
            (items[0].flag, items[0].result, items[0].unit, items[0].reference),
            ("H", "34", "mg/dL", "6~20 mg/dL"),
        )
        self.assertEqual(
            (items[1].flag, items[1].result, items[1].unit, items[1].reference),
            ("L", "133", "mmol/L", "136~145 mmol/L"),
        )
        self.assertEqual(
            (items[2].flag, items[2].result, items[2].unit, items[2].reference),
            ("", "4.3", "mmol/L", "3.5~5.1 mmol/L"),
        )
        self.assertEqual(
            (items[5].flag, items[5].result, items[5].unit, items[5].reference),
            ("H", "6.34", "mg/dL", "M:0.7~1.2; F:0.5-0.9 mg/dL"),
        )
        self.assertEqual(
            (items[6].flag, items[6].result, items[6].unit, items[6].reference),
            ("", "15", "U/L", "M:<41; F:<33 U/L"),
        )

    def test_hospital_copied_rows_with_missing_previous_result_column(self):
        text = """檢體　　：
(Specimen type)\tBlood
項目\tH/L\t結果\t單位\t參考值
Na\tL\t133\tmmol/L\t136~145 mmol/L
K\t\t4.3\t(  )\tmmol/L\t3.5~5.1 mmol/L
Creat\tH\t6.34\tmg/dL\tM:0.7~1.2; F:0.5-0.9 mg/dL
"""
        items = parse_text(text)

        self.assertEqual(
            [(item.raw_name, item.flag, item.result, item.unit, item.reference) for item in items],
            [
                ("Na", "L", "133", "mmol/L", "136~145 mmol/L"),
                ("K", "", "4.3", "mmol/L", "3.5~5.1 mmol/L"),
                ("Creat", "H", "6.34", "mg/dL", "M:0.7~1.2; F:0.5-0.9 mg/dL"),
            ],
        )

    def test_pipe_format_and_parentheses(self):
        items = parse_text("SPECIMEN: BLOOD\nCreatinine | ( 2.10 ) | mg/dL | 0.5-1.1 | H")
        self.assertEqual(items[0].result, "2.10")
        self.assertEqual(items[0].flag, "H")

    def test_inline_flag_and_specimen_heading(self):
        items = parse_text("[BLOOD]\nHb | 10.2 L | g/dL | 12.0-16.0")
        self.assertEqual(items[0].specimen, "BLOOD")
        self.assertEqual(items[0].result, "10.2")
        self.assertEqual(items[0].flag, "L")

    def test_tab_and_colon_formats(self):
        items = parse_text("SERUM\nNa\t139\tmmol/L\t136-145\nK: 4.5 mmol/L | 3.5-5.1")
        self.assertEqual(len(items), 2)
        self.assertEqual(items[0].raw_name, "Na")
        self.assertEqual(items[1].raw_name, "K")

    def test_unknown_urine_items_are_suppressed(self):
        grouped = classify_items(parse_text(
            "SPECIMEN: URINE\nUnknown Crystal | Few | | N/A"
        ))
        self.assertNotIn("驗尿檢查", grouped)

    def test_vghtpe_unmapped_assays_keep_original_fields_and_category(self):
        text = """檢體　　：
(Specimen type)\tBlood
項目\tH/L\t結果\t前次結果\t單位\t參考值
HBsAb quantitative\tH\t156.8\t(  )\tmIU/mL\tNon-reactive: <10
C-peptide\t\t2.31\t(  )\tng/mL\t1.10~4.40
Novel Tumor Index\t\t0.82\t(  )\tCOI\tNegative: <1.0
"""
        grouped = classify_items(parse_text(text))

        infection = grouped["抽血檢查"]["感染血清學檢查"][0]
        self.assertEqual(
            (infection["zh"], infection["en"], infection["result"], infection["unit"]),
            ("HBsAb quantitative", "HBsAb quantitative", "156.8", "mIU/mL"),
        )
        endocrine = grouped["抽血檢查"]["內分泌與荷爾蒙檢查"][0]
        self.assertEqual((endocrine["result"], endocrine["reference"]), ("2.31", "1.10~4.40"))
        self.assertNotIn("一般生化檢查", grouped["抽血檢查"])

    def test_duplicate_aliases_are_removed(self):
        grouped = classify_items(parse_text(
            "SPECIMEN: BLOOD\nCreatinine | 2.1 | mg/dL | 0.5-1.1\n"
            "Crea | 2.1 | mg/dL | 0.5-1.1"
        ))
        self.assertEqual(len(grouped["抽血檢查"]["腎功能檢查"]), 1)

    def test_blood_and_urine_creatinine_are_distinguished_by_specimen(self):
        text = """SPECIMEN: BLOOD
Cr | 0.72 | mg/dL | M:0.7~1.2; F:0.5-0.9
SPECIMEN: URINE(SPOT)
Cr | 98 | mg/dL |
"""
        grouped = classify_items(parse_text(text))

        blood_rows = grouped["抽血檢查"]["腎功能檢查"]
        urine_rows = grouped["驗尿檢查"]["尿蛋白與白蛋白"]
        self.assertEqual(len(blood_rows), 1)
        self.assertEqual(len(urine_rows), 1)
        self.assertEqual((blood_rows[0]["key"], blood_rows[0]["zh"]), ("Creatinine", "肌酸酐"))
        self.assertEqual((urine_rows[0]["key"], urine_rows[0]["zh"]), ("Urine Creatinine", "尿液肌酸酐"))
        self.assertEqual((blood_rows[0]["result"], urine_rows[0]["result"]), ("0.72", "98"))

    def test_urine_ratio_report_suppresses_component_noise(self):
        text = """Medical order: Crea,Protein
GLU | | |
Protein | 87.90 | mg/dL | N/A
Crea | 54.60 | mg/dL | N/A
Pro/Cr | 1.610 | | <0.2 | H
SPECIMEN: URINE(SPOT)
Medical order: Microalbumin, spot urine
Result | 75.67 | mg/dL | N/A
Total | | |
Malb/Ucre | 1.386 | | <0.03:normal;
0.03~0.3 | microalbuminurmia; | |
>0.3 | overt | proteinuria |
"""
        grouped = classify_items(parse_text(text))
        rows = [
            row
            for categories in grouped.values()
            for category_rows in categories.values()
            for row in category_rows
        ]

        self.assertEqual([row["key"] for row in rows], ["ACR", "PCR"])
        acr = next(row for row in rows if row["key"] == "ACR")
        pcr = next(row for row in rows if row["key"] == "PCR")
        self.assertEqual((acr["zh"], acr["result"], acr["reference"]), ("尿白蛋白肌酸酐比", "1.386", "<0.03:normal;"))
        self.assertEqual((pcr["zh"], pcr["result"], pcr["reference"], pcr["flag"]), ("尿蛋白肌酸酐比", "1.610", "<0.2", "H"))
        self.assertNotIn("Urine Protein", {row["key"] for row in rows})
        self.assertNotIn("Urine Creatinine", {row["key"] for row in rows})

    def test_hospital_bilirubin_and_egfr_aliases_are_categorized(self):
        grouped = classify_items(parse_text(
            "SPECIMEN: BLOOD\n"
            "T.BILI | 0.25 | mg/dL | <1.2 mg/dL\n"
            "eGFR(M) | 9 | mL/min/1.73M2 | >60 mL/min/1.73M2"
        ))

        liver_rows = grouped["抽血檢查"]["肝膽功能檢查"]
        kidney_rows = grouped["抽血檢查"]["腎功能檢查"]
        self.assertEqual([row["key"] for row in liver_rows], ["Total Bilirubin"])
        self.assertEqual([row["key"] for row in kidney_rows], ["eGFR"])
        self.assertNotIn("一般生化檢查", grouped["抽血檢查"])

    def test_single_order_result_and_separate_reference(self):
        text = """檢體　　：
(Specimen type)\tBlood
醫囑名稱：
(Medical order)\tCa++, free
H/L\t結果\t前次結果\t單位
    \t1.20\t(  )\tmmol/L

參考值:
1.13 ~ 1.31 mmol/L
檢體　　：
(Specimen type)\tBlood
醫囑名稱：
(Medical order)\t25-OH Vitamin D(SP)
H/L\t結果\t前次結果\t單位
L\t26.0\t(  )\tng/mL

參考值:
Recommendation: >=30 ng/mL
"""
        grouped = classify_items(parse_text(text))
        electrolyte = grouped["抽血檢查"]["電解質與酸鹼檢查"][0]
        vitamin = grouped["抽血檢查"]["維生素與營養檢查"][0]

        self.assertEqual(
            (electrolyte["key"], electrolyte["result"], electrolyte["unit"], electrolyte["reference"]),
            ("Ionized Calcium", "1.20", "mmol/L", "1.13 ~ 1.31 mmol/L"),
        )
        self.assertEqual(
            (vitamin["key"], vitamin["result"], vitamin["flag"], vitamin["reference"]),
            ("Vitamin D", "26.0", "L", "Recommendation: >=30 ng/mL"),
        )

    def test_single_a1c_and_duplicate_afp(self):
        text = """SPECIMEN: BLOOD
(Medical order)\tA1C (HbA1C)
H/L\t結果\t前次結果\t單位\t參考值
H\t6.1\t(  )\t%\tHealthy: 4.8~5.9%
SPECIMEN: BLOOD
(Medical order)\tAFP
H/L\t結果\t前次結果\t單位
\t3.70\t(  )\tng/mL
參考值:
<=7.0
SPECIMEN: BLOOD
(Medical order)\tAFP
H/L\t結果\t前次結果\t單位
\t3.70\t(  )\tng/mL
參考值:
<=7.0
"""
        grouped = classify_items(parse_text(text))
        a1c_rows = grouped["抽血檢查"]["血糖與糖尿病相關檢查"]
        afp_rows = grouped["抽血檢查"]["腫瘤指標"]

        self.assertEqual((a1c_rows[0]["key"], a1c_rows[0]["result"], a1c_rows[0]["flag"]), ("HbA1c", "6.1", "H"))
        self.assertEqual(len(afp_rows), 1)
        self.assertEqual((afp_rows[0]["key"], afp_rows[0]["reference"]), ("AFP", "<=7.0"))

    def test_complete_mixed_hospital_report_keeps_all_nonempty_results(self):
        text = """檢體：
(Specimen type)\tBlood
(Medical order)\tCa++, free
H/L\t結果\t前次結果\t單位
\t1.20\t(  )\tmmol/L
參考值:
1.13 ~ 1.31 mmol/L
檢體：
(Specimen type)\tBlood
(Medical order)\tCBC,DC,
項目\tH/L\t結果\t前次結果\t單位\t參考值
WBC\t\t6850\t(  )\t/uL\t4180 ~ 9380
RBC\t\t4.75\t(  )\t10^6/uL\t3.80 ~ 5.10
Hb\t\t14.2\t(  )\tg/dL\t10.9 ~ 15.6
Hct\t\t42.1\t(  )\t%\t33.4 ~ 45.5
MCV\t\t88.6\t(  )\tfL\t82.0 ~ 97.7
MCH\t\t29.9\t(  )\tpg\t27.6 ~ 33.0
MCHC\t\t33.7\t(  )\tg/dL\t32.0 ~ 35.2
RDW-CV\t\t12.5\t(  )\t%\t11.8 ~ 14.6
MPV\t\t9.4\t(  )\tfL\t9.0 ~ 12.6
Plt\t\t180000\t(  )\t/uL\t145000 ~ 383000
ANC\t\t5028\t(  )\t/uL\t1890 ~ 6760
Band\t\t0.0\t(  )\t%\t0 ~ 5
Neu.\t\t73.4\t(  )\t%\t45.0 ~ 81.2
Lym.\t\t20.0\t(  )\t%\t13.0 ~ 45.5
Mono.\t\t5.3\t(  )\t%\t3.7 ~ 8.8
Eos.\t\t0.9\t(  )\t%\t0.2 ~ 6.6
Baso.\t\t0.4\t(  )\t%\t0.1 ~ 1.6
檢體：
(Specimen type)\tBlood
(Medical order)\tNA,K,Ca,Alb,P,Crea,BUN,ALP,Mg,ALT,Bil-T,
項目\tH/L\t結果\t前次結果\t單位\t參考值
BUN\t\t12\t(  )\tmg/dL\t6~20 mg/dL
Na\t\t145\t(  )\tmmol/L\t136~145 mmol/L
K\t\t3.6\t(  )\tmmol/L\t3.5~5.1 mmol/L
Cl\t\t\t(  )\t\t
Creat\t\t0.72\t(  )\tmg/dL\tM:0.7~1.2; F:0.5-0.9 mg/dL
ALT\t\t10\t(  )\tU/L\tM:<41; F:<33 U/L
ALK-P\t\t96\t(  )\tU/L\tM:40~129 F:35-104 U/L
T.BILI\t\t0.35\t(  )\tmg/dL\t<1.2 mg/dL
I.P.\t\t2.8\t(  )\tmg/dL\t2.5~4.5 mg/dL
Ca\t\t10.0\t(  )\tmg/dL\t8.6~10.0 mg/dL
ALB\t\t4.5\t(  )\tg/dL\t3.7 ~ 5.3 g/dL
Mg\t\t2.30\t(  )\tmg/dL\t1.6~2.4 mg/dL
eGFR(M)\t\t80\t(  )\t\t>60 mL/min/1.73M^2
檢體：
(Specimen type)\tBlood
(Medical order)\t25-OH Vitamin D(SP)
H/L\t結果\t前次結果\t單位
L\t26.0\t(  )\tng/mL
參考值:
Recommendation: >=30 ng/mL
檢體：
(Specimen type)\tBlood
(Medical order)\tIntact-PTH
H/L\t結果\t前次結果\t單位
\t67.1\t(  )\tpg/mL
參考值:
17.3~74.1
檢體：
(Specimen type)\tBlood
(Medical order)\tA1C (HbA1C)
H/L\t結果\t前次結果\t單位\t參考值
H\t6.1\t(  )\t%\tHealthy: 4.8~5.9%
SPECIMEN: BLOOD
UA\t\t3.7\t(  )\tmg/dL\tM:3.4~7.0; F:2.4-5.7 mg/dL
SPECIMEN: BLOOD
CHOL.\t\t188\t(  )\tmg/dL\t<200
TG\t\t83\t(  )\tmg/dL\t<150
HDL-C\t\t52.0\t(  )\tmg/dL\tM:>40; F:>50
LDL-C(實測值)\t\t114.0\t(  )\tmg/dL\t<130
CHOL/HDLC\t\t3.6\t(  )\t\tMale 3.1~4.7, Female 2.8~4.6
SPECIMEN: BLOOD
(Medical order)\tAFP
H/L\t結果\t前次結果\t單位
\t3.70\t(  )\tng/mL
參考值:
<=7.0
SPECIMEN: BLOOD
(Medical order)\tAFP
H/L\t結果\t前次結果\t單位
\t3.70\t(  )\tng/mL
參考值:
<=7.0
"""
        grouped = classify_items(parse_text(text))
        rows = [
            row
            for categories in grouped.values()
            for category_rows in categories.values()
            for row in category_rows
        ]
        keys = {row["key"] for row in rows}

        self.assertEqual(len(rows), 40)
        self.assertTrue({
            "Ionized Calcium", "WBC", "Basophil", "BUN", "ALP",
            "Total Bilirubin", "eGFR", "Vitamin D", "Intact-PTH",
            "HbA1c", "Uric Acid", "Total Cholesterol", "LDL-C",
            "Cholesterol/HDL Ratio", "AFP",
        }.issubset(keys))
        self.assertNotIn("Chloride", keys)
        self.assertEqual(sum(row["key"] == "AFP" for row in rows), 1)


class ReportTests(unittest.TestCase):
    def test_hospital_format_report_keeps_reference_columns(self):
        grouped = classify_items(parse_text(
            "檢體　　：\n"
            "(Specimen type)\tBlood\n"
            "項目\tH/L\t結果\t前次結果\t單位\t參考值\n"
            "Cl\t\t\t(  )\t\t\n"
            "Creat\tH\t6.34\t(  )\tmg/dL\tM:0.7~1.2; F:0.5-0.9 mg/dL\n"
            "ALT\t\t15\t(  )\tU/L\tM:<41; F:<33 U/L\n"
        ))
        rows = [
            row
            for categories in grouped.values()
            for category_rows in categories.values()
            for row in category_rows
        ]

        self.assertEqual({row["key"] for row in rows}, {"Creatinine", "ALT"})
        creatinine = next(row for row in rows if row["key"] == "Creatinine")
        alt = next(row for row in rows if row["key"] == "ALT")
        self.assertEqual(creatinine["reference"], "M:0.7~1.2; F:0.5-0.9 mg/dL")
        self.assertEqual(alt["reference"], "M:<41; F:<33 U/L")

    def test_docx_columns_order_and_bold_flag(self):
        grouped = classify_items(parse_text(
            "SPECIMEN: BLOOD\nHb | 10.2 | g/dL | 12.0-16.0 | L\n"
            "SPECIMEN: URINE\nACR | 60 | mg/g | <30 | H\n"
            "Abdominal Sonography\nFinding: Test finding.\nImpression: Test impression."
        ))
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "report.docx"
            write_report(grouped, output)
            doc = Document(output)
            headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
            self.assertEqual(headings, ["1. 血液檢查", "3. 尿液檢查", "4. 影像檢查"])
            self.assertEqual(
                [cell.text for cell in doc.tables[0].rows[0].cells],
                ["中文項目", "英文項目", "結果", "正常值"],
            )
            result_cell = doc.tables[0].rows[1].cells[2]
            self.assertTrue(any(run.bold for run in result_cell.paragraphs[0].runs))
            self.assertEqual(len(doc.tables[-1].columns), 4)
            self.assertTrue(all(
                row._tr.xpath("./w:trPr/w:cantSplit")
                for table in doc.tables
                for row in table.rows
            ))
            header_alignments = [
                cell.paragraphs[0].alignment
                for cell in doc.tables[0].rows[0].cells
            ]
            body_alignments = [
                cell.paragraphs[0].alignment
                for cell in doc.tables[0].rows[1].cells
            ]
            self.assertEqual(header_alignments, [1, 1, 1, 1])
            self.assertEqual(body_alignments, [None, None, 1, 1])
            self.assertGreater(doc.sections[0].page_width, doc.sections[0].page_height)
            image_text = doc.tables[-1].rows[1].cells[2].text
            self.assertIn("Finding: Test finding.", image_text)
            self.assertIn("Impression: Test impression.", image_text)
            report_text = "\n".join(
                [paragraph.text for paragraph in doc.paragraphs]
                + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
            )
            self.assertNotIn("原始資料", report_text)
            self.assertNotIn("English", report_text)
            self.assertNotIn("其他抽血", report_text)


if __name__ == "__main__":
    unittest.main()
